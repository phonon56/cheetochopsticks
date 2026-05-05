#!/usr/bin/env python3
"""
Patch the SAP audit .docx for the DOJ Title II Interim Final Rule.

Updates every "April 24, 2026" reference to "April 26, 2027" (the new
large-public-entity deadline for jurisdictions over 50k population, which
includes Colorado Springs) and appends a short IFR-acknowledgement
paragraph at the end of the document.

Setup once:
    pip3 install python-docx

Run (auto-detects the .docx in this folder, or pass a path):
    python3 patch_audit_deadline.py
    python3 patch_audit_deadline.py SAP_Accessibility_Audit.docx
"""

from __future__ import annotations

import re
import shutil
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
except ImportError:
    sys.exit("python-docx not installed. Run: pip3 install python-docx")


HERE = Path(__file__).resolve().parent

OLD = "April 24, 2026"
NEW = "April 26, 2027"

# Map of full-text replacements (case-sensitive, exact match).
SUBS: list[tuple[str, str]] = [
    (OLD, NEW),
    ("April 24th, 2026", "April 26th, 2027"),
    ("24 April 2026", "26 April 2027"),
    ("2026-04-24", "2027-04-26"),
]

IFR_NOTE = (
    "Note on compliance deadline. The U.S. Department of Justice signed an "
    "Interim Final Rule on April 16, 2026 (effective April 20, 2026) extending "
    "the Title II web and mobile accessibility compliance deadline by one year "
    "for large public entities (population 50,000 or more), including the City "
    "of Colorado Springs. The new deadline is April 26, 2027. Smaller entities "
    "and special districts moved from April 26, 2027 to April 26, 2028. The "
    "substantive standard — WCAG 2.1 Level AA — is unchanged. Title II's "
    "underlying \"equally effective communication\" obligation also continues "
    "to apply regardless of the rulemaking date."
)


def patch_runs(paragraph) -> int:
    """
    Replace OLD -> NEW across runs in a paragraph without losing formatting.
    Handles the docx-specific case where text is split across runs.
    """
    n = 0
    full = "".join(r.text for r in paragraph.runs)
    if not any(old in full for old, _ in SUBS):
        return 0
    new_full = full
    for old, new in SUBS:
        if old in new_full:
            n += new_full.count(old)
            new_full = new_full.replace(old, new)
    if not paragraph.runs:
        return 0
    paragraph.runs[0].text = new_full
    for r in paragraph.runs[1:]:
        r.text = ""
    return n


def find_audit_doc() -> Path:
    if len(sys.argv) > 1:
        p = Path(sys.argv[1]).resolve()
        if not p.exists():
            sys.exit(f"File not found: {p}")
        return p
    candidates = sorted(HERE.glob("*[Aa]udit*.docx")) + sorted(HERE.glob("SAP*.docx"))
    candidates = [c for c in candidates if not c.name.endswith(".bak.docx")]
    if not candidates:
        sys.exit(
            f"No audit .docx found in {HERE}. "
            "Save the file here or pass the path as an argument."
        )
    if len(candidates) > 1:
        print("Multiple candidates found:")
        for c in candidates:
            print(f"  {c.name}")
        sys.exit("Pass the desired path explicitly.")
    return candidates[0]


def main() -> None:
    src = find_audit_doc()
    backup = src.with_suffix(f".bak-{date.today().isoformat()}.docx")
    shutil.copy2(src, backup)
    print(f"Backup saved: {backup.name}")

    doc = Document(src)
    total = 0

    for p in doc.paragraphs:
        total += patch_runs(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total += patch_runs(p)

    for section in doc.sections:
        for hf in (section.header, section.footer, section.first_page_header,
                   section.first_page_footer, section.even_page_header,
                   section.even_page_footer):
            for p in hf.paragraphs:
                total += patch_runs(p)

    full = "\n".join(p.text for p in doc.paragraphs)
    if "Interim Final Rule" not in full:
        doc.add_paragraph()
        h = doc.add_paragraph()
        run = h.add_run("Compliance deadline update")
        run.bold = True
        doc.add_paragraph(IFR_NOTE)
        print("Appended IFR acknowledgement paragraph.")
    else:
        print("IFR acknowledgement already present; skipped append.")

    doc.save(src)
    print(f"Patched {total} date occurrence(s) in {src.name}.")
    print(f"  {OLD!r} -> {NEW!r}")


if __name__ == "__main__":
    main()
